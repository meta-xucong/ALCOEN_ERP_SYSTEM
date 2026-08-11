"""Business services for the independent official-contract generator."""

from __future__ import annotations

import hashlib
import json
import os
import re
from copy import deepcopy
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Iterable, Optional
from uuid import uuid4

from docx import Document as WordDocument
from docxtpl import DocxTemplate
from flask import current_app
from sqlalchemy import or_
from werkzeug.utils import secure_filename

from app import db
from app.models import (
    Company,
    FormalContract,
    FormalContractDocument,
    FormalContractItem,
    FormalContractParty,
    FormalContractSync,
    FormalContractTemplate,
    Department,
    Product,
)
from app.services.contract_service import ContractService
from app.services.product_service import ProductService


DEFAULT_PARTY_B = {
    'name': '江苏纯安科技有限公司',
    'billing_address': '南京市江北新区新锦湖路3-1号中丹生态生命科学产业园1期A座633室',
    'phone': '13813809690',
    'tax_no': '91320113MA1YQB4L97',
    'bank_name': '南京银行南京金融城支行',
    'bank_account': '0162210000004639',
}

TEMPLATE_ROOTS = {'contract', 'party_a', 'party_b', 'item'}
TEMPLATE_TOKEN_RE = re.compile(r'{{\s*([A-Za-z_][\w.]*)\s*}}')


class OfficialContractValidationError(ValueError):
    """Raised when an official-contract payload cannot be saved or generated."""


def _text(value: Any) -> str:
    return str(value or '').strip()


def _money(value: Any, default: float = 0.0) -> float:
    try:
        return float(
            Decimal(str(value if value not in (None, '') else default)).quantize(
                Decimal('0.01'),
                rounding=ROUND_HALF_UP,
            )
        )
    except Exception:
        return default


def _date_value(value: Any) -> Optional[date]:
    if isinstance(value, date):
        return value
    if not value:
        return None
    try:
        return datetime.strptime(str(value), '%Y-%m-%d').date()
    except ValueError as exc:
        raise OfficialContractValidationError(
            '签订日期格式不正确，请使用 YYYY-MM-DD'
        ) from exc


def _format_date(value: Optional[date]) -> str:
    return value.strftime('%Y-%m-%d') if value else ''


def _format_contract_sign_date(value: Any) -> str:
    """Format a signature date exactly as the supplied DOCX template does."""
    text = _format_date(value) if isinstance(value, date) else _text(value)
    match = re.fullmatch(r'(\d{4})-(\d{2})-(\d{2})', text)
    if match:
        return f'{match.group(1)}年{match.group(2)}月{match.group(3)}日'
    return text


def _section_to_upper(section: int) -> str:
    digits = '零壹贰叁肆伍陆柒捌玖'
    units = ('', '拾', '佰', '仟')
    if section == 0:
        return ''

    result: list[str] = []
    zero_pending = False
    for index in range(4):
        digit = section % 10
        section //= 10
        if digit:
            if zero_pending and result:
                result.append('零')
            result.append(digits[digit] + units[index])
            zero_pending = False
        elif result:
            zero_pending = True
        if section == 0:
            break
    return ''.join(reversed(result))


def amount_to_upper(value: Any) -> str:
    """Convert a monetary value to a Chinese RMB uppercase string."""
    amount = Decimal(str(value or 0)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    if amount < 0:
        return f'负{amount_to_upper(-amount)}'

    integer_part = int(amount)
    fraction = int((amount - integer_part) * 100)
    digits = '零壹贰叁肆伍陆柒捌玖'
    if integer_part == 0:
        integer_text = '零'
    else:
        sections: list[str] = []
        remaining = integer_part
        while remaining:
            sections.append(remaining % 10000)
            remaining //= 10000
        big_units = ('', '万', '亿', '兆')
        integer_parts: list[str] = []
        zero_between = False
        for index in range(len(sections) - 1, -1, -1):
            section = sections[index]
            if section == 0:
                if integer_parts:
                    zero_between = True
                continue
            section_text = _section_to_upper(section)
            if zero_between and integer_parts and not integer_parts[-1].endswith('零'):
                integer_parts.append('零')
            integer_parts.append(section_text + big_units[index])
            zero_between = section < 1000
        integer_text = ''.join(integer_parts)

    jiao = fraction // 10
    fen = fraction % 10
    fraction_text = ''
    if jiao:
        fraction_text += digits[jiao] + '角'
    if fen:
        fraction_text += digits[fen] + '分'
    if not fraction_text:
        fraction_text = '整'
    return f'人民币{integer_text}元{fraction_text}'


class OfficialContractService:
    """Persistence, template, and synchronization logic for formal contracts."""

    @staticmethod
    def _set_paragraph_text(paragraph, value: Any) -> None:
        """Replace paragraph text while retaining the first run's formatting."""
        text = str(value or '')
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ''
        else:
            paragraph.add_run(text)

    @staticmethod
    def _set_cell_text(cell, value: Any) -> None:
        """Replace a table cell with plain text while retaining the cell style."""
        # Assigning ``cell.text`` recreates the cell paragraphs and discards
        # the paragraph/run formatting from the supplied DOCX template.
        lines = str(value or '').splitlines() or ['']
        paragraphs = cell.paragraphs or [cell.add_paragraph()]
        for index, line in enumerate(lines):
            if index >= len(paragraphs):
                paragraph = cell.add_paragraph()
                if paragraphs[-1]._p.pPr is not None:
                    paragraph._p.insert(0, deepcopy(paragraphs[-1]._p.pPr))
                paragraphs.append(paragraph)
            OfficialContractService._set_paragraph_text(paragraphs[index], line)
        for paragraph in paragraphs[len(lines):]:
            OfficialContractService._set_paragraph_text(paragraph, '')

    @staticmethod
    def _replace_paragraph_label(document, label: str, value: Any) -> bool:
        """Replace the content of the first paragraph beginning with a label."""
        for paragraph in document.paragraphs:
            if paragraph.text.lstrip().startswith(label):
                OfficialContractService._set_paragraph_text(
                    paragraph,
                    f'{label}{value or ""}',
                )
                return True
        return False

    @staticmethod
    def _is_fixed_layout_template(document) -> bool:
        """Recognize the supplied static contract template layout."""
        all_text = '\n'.join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        return (
            len(document.tables) >= 3
            and '合同编号' in all_text
            and '品名' in all_text
            and '需      方' in all_text
        )

    @staticmethod
    def _render_fixed_layout_template(document, context: dict) -> None:
        """Fill the fixed-layout DOCX template supplied for accessory contracts."""
        contract = context['contract']
        party_a = context['party_a']
        party_b = context['party_b']
        items = context['items']

        if len(document.tables) < 3:
            raise OfficialContractValidationError('固定版式合同模板缺少必要表格')

        meta_table = document.tables[0]
        meta_values = (
            contract['contract_no'],
            contract['sign_place'],
            contract.get(
                'sign_date_display',
                _format_contract_sign_date(contract.get('sign_date')),
            ),
        )
        for row, value in zip(meta_table.rows, meta_values):
            if len(row.cells) > 1:
                OfficialContractService._set_cell_text(row.cells[1], value)

        items_table = document.tables[1]
        template_row = items_table.rows[1]._tr
        total_row = items_table.rows[-1]._tr
        for _ in range(max(0, len(items) - 1)):
            total_row.addprevious(deepcopy(template_row))

        item_rows = items_table.rows[1:1 + len(items)]
        for index, (row, item) in enumerate(zip(item_rows, items), start=1):
            values = (
                index,
                item['product_name'] or item['product_code'],
                item['product_model'],
                item['quantity_display'],
                f"￥{item['unit_price_display']}",
                f"￥{item['total_amount_display']}",
            )
            for cell, value in zip(row.cells, values):
                OfficialContractService._set_cell_text(cell, value)

        total_cells = items_table.rows[-1].cells
        if len(total_cells) >= 6:
            OfficialContractService._set_cell_text(
                total_cells[1],
                f"合计：{contract['total_amount_upper']}",
            )
            OfficialContractService._set_cell_text(
                total_cells[5],
                f"￥{contract['total_amount_display']}",
            )

        paragraph_values = (
            ('需方（甲方）：', party_a['name']),
            ('供方（乙方）：', party_b['name']),
            ('质量标准：', contract['quality_standard']),
            ('交货方式及费用承担：', contract['delivery_terms']),
            ('交货时间及地点：', contract['delivery_schedule']),
            ('结算方式及期限：', contract['settlement_terms']),
            ('违约责任：', contract['breach_terms']),
            ('解决合同纠纷的方式：', contract['dispute_terms']),
        )
        for label, value in paragraph_values:
            OfficialContractService._replace_paragraph_label(document, label, value)

        party_table = document.tables[2]
        if party_table.rows and len(party_table.rows[0].cells) >= 2:
            party_a_text = (
                '需      方 (甲方)\n\n'
                f"公司名称：{party_a['name']}\n"
                f"开票地址：{party_a['billing_address']}\n"
                f"公司电话：{party_a['phone']}\n"
                f"税    号：{party_a['tax_no']}\n"
                f"开户银行：{party_a['bank_name']}\n"
                f"银行账号：{party_a['bank_account']}"
            )
            party_b_text = (
                '供     方（乙方）\n\n'
                f"公司名称：{party_b['name']}\n"
                f"开票地址：{party_b['billing_address']}\n"
                f"公司电话：{party_b['phone']}\n"
                f"税    号：{party_b['tax_no']}\n"
                f"开户银行：{party_b['bank_name']}\n"
                f"银行账号：{party_b['bank_account']}"
            )
            OfficialContractService._set_cell_text(
                party_table.rows[0].cells[0],
                party_a_text,
            )
            OfficialContractService._set_cell_text(
                party_table.rows[0].cells[1],
                party_b_text,
            )

    @staticmethod
    def _template_folder() -> Path:
        configured = current_app.config.get('OFFICIAL_CONTRACT_TEMPLATE_FOLDER')
        folder = Path(configured or (Path(current_app.root_path).parent / 'data' / 'official_contract_templates'))
        folder.mkdir(parents=True, exist_ok=True)
        return folder

    @staticmethod
    def _builtin_template_source() -> Path:
        """Return the DOCX bundled with the application distribution."""
        return (
            Path(current_app.root_path).parent
            / 'templates'
            / 'official_contract'
            / 'builtin'
            / '配件合同模版.docx'
        )

    @staticmethod
    def ensure_builtin_template() -> FormalContractTemplate | None:
        """Register the bundled template once and keep it as global fallback."""
        source = OfficialContractService._builtin_template_source()
        if not source.is_file():
            return None

        active_global = FormalContractTemplate.query.filter_by(
            status='active',
            department_id=None,
        ).order_by(
            FormalContractTemplate.activated_at.desc(),
            FormalContractTemplate.id.desc(),
        ).first()
        if active_global:
            return active_global

        file_hash = hashlib.sha256(source.read_bytes()).hexdigest()
        template = FormalContractTemplate.query.filter_by(
            file_hash=file_hash,
            department_id=None,
        ).first()
        if not template:
            stored_path = OfficialContractService._template_folder() / (
                f'builtin_{file_hash[:16]}.docx'
            )
            if not stored_path.is_file():
                stored_path.write_bytes(source.read_bytes())
            validation = OfficialContractService.validate_template_file(stored_path)
            if validation['unknown'] or validation['mode'] == 'unsupported_static':
                raise OfficialContractValidationError(
                    '系统内置正式合同模板无法通过模板校验'
                )
            template = FormalContractTemplate(
                department_id=None,
                name='配件合同默认模板',
                version='builtin-v1',
                original_filename=source.name,
                stored_path=str(stored_path),
                file_hash=file_hash,
                status='inactive',
                description='系统内置默认模板。管理员可按部门上传并启用新的 DOCX 模板。',
                uploaded_by_id=None,
            )
            db.session.add(template)

        template.status = 'active'
        template.activated_at = template.activated_at or datetime.now()
        template.deactivated_at = None
        db.session.commit()
        return template

    @staticmethod
    def _export_folder() -> Path:
        configured = current_app.config.get('OFFICIAL_CONTRACT_EXPORT_FOLDER')
        folder = Path(configured or (Path(current_app.root_path).parent / 'exports' / 'official_contracts'))
        folder.mkdir(parents=True, exist_ok=True)
        return folder

    @staticmethod
    def _normalize_party_name(name: Any) -> str:
        return re.sub(r'\s+', ' ', _text(name))

    @staticmethod
    def _normalize_items(items: Iterable[dict]) -> list[dict]:
        normalized: list[dict] = []
        for index, raw in enumerate(items):
            raw = raw or {}
            product = None
            product_id = raw.get('product_id')
            if product_id not in (None, ''):
                try:
                    product = Product.query.get(int(product_id))
                except (TypeError, ValueError):
                    product = None
                if not product:
                    raise OfficialContractValidationError('选择的产品不存在')

            product_code = _text(raw.get('product_code')) or _text(
                getattr(product, 'product_code', None)
            )
            product_name = _text(raw.get('product_name')) or _text(
                getattr(product, 'product_name', None)
            )
            product_model = _text(raw.get('product_model')) or _text(
                getattr(product, 'product_model', None)
            )
            unit = _text(raw.get('unit')) or '个'
            quantity = _money(raw.get('quantity'))
            unit_price = _money(raw.get('unit_price', raw.get('price')))
            if not product_code and not product_name:
                continue
            if quantity <= 0:
                raise OfficialContractValidationError(f'第 {index + 1} 行产品数量必须大于 0')
            if not product_code:
                raise OfficialContractValidationError(
                    f'第 {index + 1} 行缺少产品编码，无法同步到交易合同'
                )
            normalized.append({
                'product_id': product.id if product else None,
                'product_code': product_code,
                'product_name': product_name,
                'product_model': product_model,
                'unit': unit,
                'quantity': quantity,
                'unit_price': unit_price,
                'total_amount': _money(quantity * unit_price),
                'remark': _text(raw.get('remark')),
                'sort_order': index,
            })

        if not normalized:
            raise OfficialContractValidationError('至少添加一条产品明细')
        return normalized

    @staticmethod
    def _find_or_create_party(data: dict) -> FormalContractParty:
        party_name = OfficialContractService._normalize_party_name(data.get('party_a_name'))
        if not party_name:
            raise OfficialContractValidationError('甲方名称不能为空')

        company = Company.query.filter_by(name=party_name).first()
        company_exists = company is not None
        party = FormalContractParty.query.filter_by(party_a_name=party_name).first()
        if not company:
            company = Company(name=party_name)
            db.session.add(company)
            db.session.flush()
        if not party:
            party = FormalContractParty(
                party_a_name=party_name,
                company_id=company.id,
                source='erp_company' if company_exists else 'manual',
            )
            db.session.add(party)
        elif party.company_id is None:
            party.company_id = company.id

        for field in ('billing_address', 'phone', 'tax_no', 'bank_name', 'bank_account'):
            if field in data:
                setattr(party, field, _text(data.get(field)) or None)
        party.last_used_at = datetime.now()
        # The formal contract stores the party foreign key directly. Flush the
        # new party before the caller assigns party.id to the contract.
        db.session.flush()
        return party

    @staticmethod
    def get_party_defaults(party_id: int | None = None, party_name: str | None = None) -> dict:
        party = None
        if party_id:
            party = FormalContractParty.query.get(party_id)
        elif party_name:
            normalized_name = OfficialContractService._normalize_party_name(party_name)
            party = FormalContractParty.query.filter_by(party_a_name=normalized_name).first()

        if not party and party_name:
            company = Company.query.filter_by(
                name=OfficialContractService._normalize_party_name(party_name)
            ).first()
            if company:
                return {
                    'party_id': None,
                    'company_id': company.id,
                    'party_a_name': company.name,
                    'billing_address': '',
                    'phone': '',
                    'tax_no': '',
                    'bank_name': '',
                    'bank_account': '',
                    **OfficialContractService._latest_contract_defaults(None),
                }
            return {
                'party_id': None,
                'company_id': None,
                'party_a_name': OfficialContractService._normalize_party_name(party_name),
                'billing_address': '',
                'phone': '',
                'tax_no': '',
                'bank_name': '',
                'bank_account': '',
                **OfficialContractService._latest_contract_defaults(None),
            }

        if not party:
            return {
                'party_id': None,
                'company_id': None,
                'party_a_name': '',
                'billing_address': '',
                'phone': '',
                'tax_no': '',
                'bank_name': '',
                'bank_account': '',
                **OfficialContractService._latest_contract_defaults(None),
            }

        latest = party.formal_contracts[0] if party.formal_contracts else None
        defaults = OfficialContractService._latest_contract_defaults(latest)
        return {
            'party_id': party.id,
            'company_id': party.company_id,
            'party_a_name': party.party_a_name,
            'billing_address': party.billing_address or '',
            'phone': party.phone or '',
            'tax_no': party.tax_no or '',
            'bank_name': party.bank_name or '',
            'bank_account': party.bank_account or '',
            **defaults,
        }

    @staticmethod
    def _latest_contract_defaults(contract: FormalContract | None) -> dict:
        if not contract:
            return {
                'party_b_name': DEFAULT_PARTY_B['name'],
                'party_b_billing_address': DEFAULT_PARTY_B['billing_address'],
                'party_b_phone': DEFAULT_PARTY_B['phone'],
                'party_b_tax_no': DEFAULT_PARTY_B['tax_no'],
                'party_b_bank_name': DEFAULT_PARTY_B['bank_name'],
                'party_b_bank_account': DEFAULT_PARTY_B['bank_account'],
                'contract_no': '',
                'sign_place': '南京',
                'sign_date': '',
                'quality_standard': '符合厂家质量标准。如有疑问，双方协商解决。',
                'delivery_terms': '快递到需方，供方承担运费。',
                'delivery_schedule': '现货7个工作日，快递到需方指定地址。',
                'settlement_terms': (
                    '需方在合同签订三天内将货款全额汇入供方账户，供方在收到款项后发货，'
                    '并在交付货物后无异议，开具13%增值税专用发票交需方。'
                ),
                'breach_terms': '以《中华人民共和国民法典》为准。',
                'dispute_terms': '以《中华人民共和国民法典》为准。',
            }
        return {
            'party_b_name': contract.party_b_name or DEFAULT_PARTY_B['name'],
            'party_b_billing_address': contract.party_b_billing_address or '',
            'party_b_phone': contract.party_b_phone or '',
            'party_b_tax_no': contract.party_b_tax_no or '',
            'party_b_bank_name': contract.party_b_bank_name or '',
            'party_b_bank_account': contract.party_b_bank_account or '',
            'contract_no': contract.contract_no or '',
            'sign_place': contract.sign_place or '',
            'sign_date': _format_date(contract.sign_date),
            'quality_standard': contract.quality_standard or '',
            'delivery_terms': contract.delivery_terms or '',
            'delivery_schedule': contract.delivery_schedule or '',
            'settlement_terms': contract.settlement_terms or '',
            'breach_terms': contract.breach_terms or '',
            'dispute_terms': contract.dispute_terms or '',
        }

    @staticmethod
    def search_parties(keyword: str = '', limit: int = 20) -> list[dict]:
        keyword = _text(keyword)
        party_query = FormalContractParty.query
        company_query = Company.query
        if keyword:
            pattern = f'%{keyword}%'
            party_query = party_query.filter(FormalContractParty.party_a_name.ilike(pattern))
            company_query = company_query.filter(Company.name.ilike(pattern))

        results: list[dict] = []
        seen: set[str] = set()
        for party in party_query.order_by(
            FormalContractParty.last_used_at.desc(),
            FormalContractParty.party_a_name.asc(),
        ).limit(limit).all():
            seen.add(party.party_a_name)
            results.append({
                'party_id': party.id,
                'company_id': party.company_id,
                'name': party.party_a_name,
                'source': 'formal_contract',
            })

        for company in company_query.order_by(Company.name.asc()).limit(limit).all():
            if company.name in seen:
                continue
            results.append({
                'party_id': None,
                'company_id': company.id,
                'name': company.name,
                'source': 'erp_company',
            })
            seen.add(company.name)
            if len(results) >= limit:
                break
        return results[:limit]

    @staticmethod
    def save_formal_contract(
        data: dict,
        items: Iterable[dict],
        user_id: int | None,
        formal_contract_id: int | None = None,
    ) -> FormalContract:
        normalized_items = OfficialContractService._normalize_items(items)
        party = OfficialContractService._find_or_create_party(data)
        department_id = data.get('department_id')
        if department_id in (None, ''):
            department_id = None
        else:
            try:
                department_id = int(department_id)
            except (TypeError, ValueError) as exc:
                raise OfficialContractValidationError('合同所属部门无效') from exc
            if not Department.query.get(department_id):
                raise OfficialContractValidationError('合同所属部门不存在')

        if formal_contract_id:
            formal_contract = FormalContract.query.get(formal_contract_id)
            if not formal_contract:
                raise OfficialContractValidationError('正式合同不存在')
            if formal_contract.is_synced:
                raise OfficialContractValidationError('已同步交易合同的正式合同不能直接编辑')
        else:
            formal_contract = FormalContract(
                party_id=party.id,
                created_by_id=user_id,
            )
            db.session.add(formal_contract)

        formal_contract.party_id = party.id
        formal_contract.department_id = department_id
        formal_contract.party_a_billing_address = _text(data.get('billing_address')) or None
        formal_contract.party_a_phone = _text(data.get('phone')) or None
        formal_contract.party_a_tax_no = _text(data.get('tax_no')) or None
        formal_contract.party_a_bank_name = _text(data.get('bank_name')) or None
        formal_contract.party_a_bank_account = _text(data.get('bank_account')) or None
        formal_contract.party_b_name = _text(data.get('party_b_name')) or DEFAULT_PARTY_B['name']
        formal_contract.party_b_billing_address = _text(
            data.get('party_b_billing_address')
        ) or DEFAULT_PARTY_B['billing_address']
        formal_contract.party_b_phone = _text(data.get('party_b_phone')) or DEFAULT_PARTY_B['phone']
        formal_contract.party_b_tax_no = _text(data.get('party_b_tax_no')) or DEFAULT_PARTY_B['tax_no']
        formal_contract.party_b_bank_name = _text(
            data.get('party_b_bank_name')
        ) or DEFAULT_PARTY_B['bank_name']
        formal_contract.party_b_bank_account = _text(
            data.get('party_b_bank_account')
        ) or DEFAULT_PARTY_B['bank_account']
        formal_contract.contract_no = _text(data.get('contract_no')) or None
        formal_contract.sign_place = _text(data.get('sign_place')) or None
        formal_contract.sign_date = _date_value(data.get('sign_date'))
        for field in (
            'quality_standard',
            'delivery_terms',
            'delivery_schedule',
            'settlement_terms',
            'breach_terms',
            'dispute_terms',
        ):
            setattr(formal_contract, field, _text(data.get(field)) or None)
        formal_contract.total_amount = _money(
            sum(item['total_amount'] for item in normalized_items)
        )
        formal_contract.total_amount_upper = amount_to_upper(formal_contract.total_amount)
        formal_contract.status = 'draft'

        formal_contract.items.clear()
        for item in normalized_items:
            formal_contract.items.append(FormalContractItem(**item))

        db.session.commit()
        return formal_contract

    @staticmethod
    def build_context(formal_contract: FormalContract) -> dict:
        party = formal_contract.party
        contract_items = [
            {
                'id': item.id,
                'product_id': item.product_id,
                'product_code': item.product_code or '',
                'product_name': item.product_name or '',
                'product_model': item.product_model or '',
                'unit': item.unit or '个',
                'quantity': item.quantity,
                'unit_price': item.unit_price,
                'total_amount': item.total_amount,
                'quantity_display': f'{item.quantity:g}',
                'unit_price_display': f'{item.unit_price:.2f}',
                'total_amount_display': f'{item.total_amount:.2f}',
                'remark': item.remark or '',
            }
            for item in formal_contract.items
        ]
        party_a = {
            'name': party.party_a_name,
            'billing_address': (
                formal_contract.party_a_billing_address
                if formal_contract.party_a_billing_address is not None
                else party.billing_address
                or ''
            ),
            'phone': (
                formal_contract.party_a_phone
                if formal_contract.party_a_phone is not None
                else party.phone
                or ''
            ),
            'tax_no': (
                formal_contract.party_a_tax_no
                if formal_contract.party_a_tax_no is not None
                else party.tax_no
                or ''
            ),
            'bank_name': (
                formal_contract.party_a_bank_name
                if formal_contract.party_a_bank_name is not None
                else party.bank_name
                or ''
            ),
            'bank_account': (
                formal_contract.party_a_bank_account
                if formal_contract.party_a_bank_account is not None
                else party.bank_account
                or ''
            ),
        }
        party_b = {
            'name': formal_contract.party_b_name or '',
            'billing_address': formal_contract.party_b_billing_address or '',
            'phone': formal_contract.party_b_phone or '',
            'tax_no': formal_contract.party_b_tax_no or '',
            'bank_name': formal_contract.party_b_bank_name or '',
            'bank_account': formal_contract.party_b_bank_account or '',
        }
        contract = {
            'id': formal_contract.id,
            'contract_no': formal_contract.contract_no or '',
            'sign_place': formal_contract.sign_place or '',
            'sign_date': _format_date(formal_contract.sign_date),
            'sign_date_display': _format_contract_sign_date(
                formal_contract.sign_date
            ),
            'party_a_name': party_a['name'],
            'party_b_name': party_b['name'],
            'quality_standard': formal_contract.quality_standard or '',
            'delivery_terms': formal_contract.delivery_terms or '',
            'delivery_schedule': formal_contract.delivery_schedule or '',
            'settlement_terms': formal_contract.settlement_terms or '',
            'breach_terms': formal_contract.breach_terms or '',
            'dispute_terms': formal_contract.dispute_terms or '',
            'total_amount': formal_contract.total_amount,
            'total_amount_display': f'{formal_contract.total_amount:.2f}',
            'total_amount_upper': formal_contract.total_amount_upper or amount_to_upper(
                formal_contract.total_amount
            ),
            'items': contract_items,
        }
        return {
            'contract': contract,
            'party_a': party_a,
            'party_b': party_b,
            'items': contract_items,
            # Keep a convenient single-item alias for simple templates that do
            # not use a Jinja loop around the product table.
            'item': contract_items[0] if contract_items else {},
        }

    @staticmethod
    def validate_template_file(file_path: str | os.PathLike) -> dict:
        """Validate a DOCX template and return its detected root variables."""
        try:
            template = DocxTemplate(str(file_path))
            variables = set(template.get_undeclared_template_variables())
        except Exception as exc:
            raise OfficialContractValidationError(
                f'模板不是有效的 DOCX 文件：{exc}'
            ) from exc

        document = WordDocument(str(file_path))
        mode = 'placeholder'
        if not variables:
            mode = (
                'fixed_layout'
                if OfficialContractService._is_fixed_layout_template(document)
                else 'unsupported_static'
            )
        unknown = sorted(variables - TEMPLATE_ROOTS)
        warnings = []
        if mode == 'fixed_layout':
            warnings.append('未检测到占位符，将按标准固定版式填充合同字段。')
        elif mode == 'unsupported_static':
            warnings.append('模板没有检测到占位符，且不符合标准固定版式。')
        return {
            'variables': sorted(variables),
            'unknown': unknown,
            'mode': mode,
            'warnings': warnings,
        }

    @staticmethod
    def create_template(
        file_storage,
        *,
        name: str,
        version: str,
        description: str | None,
        uploaded_by_id: int | None,
        department_id: int | None = None,
    ) -> FormalContractTemplate:
        if department_id is not None and not Department.query.get(department_id):
            raise OfficialContractValidationError('模板所属部门不存在')
        original_filename = _text(getattr(file_storage, 'filename', ''))
        if not file_storage or not original_filename:
            raise OfficialContractValidationError('请选择模板文件')
        if not original_filename.lower().endswith('.docx'):
            raise OfficialContractValidationError('正式合同模板只支持 DOCX 格式')

        safe_name = secure_filename(original_filename) or f'template_{uuid4().hex}.docx'
        if not safe_name.lower().endswith('.docx'):
            safe_name += '.docx'
        stored_name = f'{uuid4().hex}_{safe_name}'
        stored_path = OfficialContractService._template_folder() / stored_name
        file_storage.save(stored_path)

        try:
            validation = OfficialContractService.validate_template_file(stored_path)
        except Exception:
            stored_path.unlink(missing_ok=True)
            raise

        file_hash = hashlib.sha256(stored_path.read_bytes()).hexdigest()
        template = FormalContractTemplate(
            department_id=department_id,
            name=_text(name) or Path(original_filename).stem,
            version=_text(version) or datetime.now().strftime('%Y%m%d%H%M%S'),
            original_filename=original_filename,
            stored_path=str(stored_path),
            file_hash=file_hash,
            status='inactive',
            description=_text(description) or None,
            uploaded_by_id=uploaded_by_id,
        )
        db.session.add(template)
        db.session.commit()
        template.validation = validation
        return template

    @staticmethod
    def activate_template(template_id: int) -> FormalContractTemplate:
        template = FormalContractTemplate.query.get(template_id)
        if not template:
            raise OfficialContractValidationError('模板不存在')
        validation = OfficialContractService.validate_template_file(template.stored_path)
        if validation['unknown']:
            raise OfficialContractValidationError(
                f'模板包含未知字段：{", ".join(validation["unknown"])}'
            )
        if validation['mode'] == 'unsupported_static':
            raise OfficialContractValidationError(
                '模板没有可识别的合同字段，请加入占位符或使用标准固定版式模板'
            )
        now = datetime.now()
        FormalContractTemplate.query.filter(
            FormalContractTemplate.status == 'active',
            FormalContractTemplate.department_id == template.department_id,
            FormalContractTemplate.id != template.id,
        ).update(
            {
                FormalContractTemplate.status: 'inactive',
                FormalContractTemplate.deactivated_at: now,
            },
            synchronize_session=False,
        )
        template.status = 'active'
        template.activated_at = now
        template.deactivated_at = None
        db.session.commit()
        return template

    @staticmethod
    def get_active_template(
        department_id: int | None = None,
    ) -> FormalContractTemplate | None:
        """Get a department template first, then the active global fallback."""
        if department_id is not None:
            department_template = FormalContractTemplate.query.filter_by(
                status='active',
                department_id=department_id,
            ).order_by(
                FormalContractTemplate.activated_at.desc(),
                FormalContractTemplate.id.desc(),
            ).first()
            if department_template:
                return department_template

        template = FormalContractTemplate.query.filter_by(
            status='active',
            department_id=None,
        ).order_by(
            FormalContractTemplate.activated_at.desc(),
            FormalContractTemplate.id.desc(),
        ).first()
        if template:
            return template

        return OfficialContractService.ensure_builtin_template()

    @staticmethod
    def generate_document(
        formal_contract_id: int,
        *,
        user_id: int | None,
        template_id: int | None = None,
    ) -> FormalContractDocument:
        formal_contract = FormalContract.query.get(formal_contract_id)
        if not formal_contract:
            raise OfficialContractValidationError('正式合同不存在')
        if not formal_contract.contract_no:
            raise OfficialContractValidationError('请先填写合同编号')
        if not formal_contract.items:
            raise OfficialContractValidationError('请先添加产品明细')

        template = (
            FormalContractTemplate.query.get(template_id)
            if template_id else OfficialContractService.get_active_template(
                formal_contract.department_id,
            )
        )
        if not template or not template.is_active:
            raise OfficialContractValidationError('暂无启用的正式合同模板')
        validation = OfficialContractService.validate_template_file(template.stored_path)
        if validation['unknown']:
            raise OfficialContractValidationError(
                f'模板包含未知字段：{", ".join(validation["unknown"])}'
            )

        context = OfficialContractService.build_context(formal_contract)
        if validation['mode'] == 'unsupported_static':
            raise OfficialContractValidationError(
                '当前模板没有可识别的合同字段，请加入占位符或使用标准固定版式模板'
            )
        if validation['mode'] == 'fixed_layout':
            doc = WordDocument(template.stored_path)
            OfficialContractService._render_fixed_layout_template(doc, context)
        else:
            doc = DocxTemplate(template.stored_path)
            doc.render(context, autoescape=True)

        timestamp = datetime.now().strftime('%Y%m%d%H%M%S%f')
        output_path = OfficialContractService._export_folder() / (
            f'formal_contract_{formal_contract.id}_{timestamp}.docx'
        )
        doc.save(output_path)
        file_hash = hashlib.sha256(output_path.read_bytes()).hexdigest()
        document = FormalContractDocument(
            formal_contract_id=formal_contract.id,
            template_id=template.id,
            template_version=template.version,
            docx_path=str(output_path),
            snapshot_json=json.dumps(context, ensure_ascii=False, default=str),
            file_hash=file_hash,
            generated_by_id=user_id,
        )
        db.session.add(document)
        formal_contract.status = 'generated'
        db.session.commit()
        return document

    @staticmethod
    def get_document_context(
        document: FormalContractDocument,
        *,
        fallback_formal_contract: FormalContract | None = None,
    ) -> dict:
        """Return the immutable context saved with a generated DOCX file.

        Old records created before snapshots were introduced remain printable
        through the supplied formal-contract fallback.
        """
        try:
            snapshot = json.loads(document.snapshot_json or '')
            if isinstance(snapshot, dict):
                contract = snapshot.get('contract')
                if isinstance(contract, dict):
                    contract.setdefault(
                        'sign_date_display',
                        _format_contract_sign_date(contract.get('sign_date')),
                    )
                return snapshot
        except (TypeError, json.JSONDecodeError):
            pass

        formal_contract = fallback_formal_contract or document.formal_contract
        if not formal_contract:
            raise OfficialContractValidationError(
                '正式合同文件缺少可用于打印的数据快照'
            )
        return OfficialContractService.build_context(formal_contract)

    @staticmethod
    def mark_printed(document_id: int) -> FormalContractDocument:
        document = FormalContractDocument.query.get(document_id)
        if not document:
            raise OfficialContractValidationError('正式合同文件不存在')
        document.print_count += 1
        document.last_printed_at = datetime.now()
        db.session.commit()
        return document

    @staticmethod
    def sync_to_transaction_contract(
        formal_contract_id: int,
        *,
        user_id: int | None,
    ) -> FormalContractSync:
        formal_contract = FormalContract.query.get(formal_contract_id)
        if not formal_contract:
            raise OfficialContractValidationError('正式合同不存在')
        existing_link = FormalContractSync.query.filter_by(
            formal_contract_id=formal_contract.id,
            sync_status='success',
        ).first()
        if existing_link:
            if formal_contract.status != 'synced':
                formal_contract.status = 'synced'
                db.session.commit()
            return existing_link
        if not formal_contract.contract_no:
            raise OfficialContractValidationError('请先填写合同编号')
        if not formal_contract.latest_document:
            raise OfficialContractValidationError('请先生成正式合同文件')

        party = formal_contract.party
        if not party.company_id:
            raise OfficialContractValidationError('甲方尚未关联 ERP 公司')

        products_data = []
        for item in formal_contract.items:
            if not item.product_code:
                raise OfficialContractValidationError('产品编码不能为空')
            products_data.append({
                'product_code': item.product_code,
                'product_name': item.product_name,
                'product_model': item.product_model,
                'product_type': item.product.product_type if item.product else None,
                'quantity': item.quantity,
                'unit': item.unit,
                'price': item.unit_price,
                'total': item.total_amount,
                'remark': item.remark,
            })

        try:
            contract = ContractService.create_contract(
                {
                    'contract_no': formal_contract.contract_no,
                    'company_name': party.party_a_name,
                    'created_by_id': user_id,
                    'total_value': formal_contract.total_amount,
                    'actual_received_value': formal_contract.total_amount,
                },
                products_data,
                auto_commit=False,
            )
            sync = FormalContractSync(
                formal_contract_id=formal_contract.id,
                contract_id=contract.id,
                synced_by_id=user_id,
                sync_status='success',
            )
            db.session.add(sync)
            formal_contract.status = 'synced'
            db.session.commit()
            return sync
        except Exception as exc:
            db.session.rollback()
            raise OfficialContractValidationError(f'同步交易合同失败：{exc}') from exc

    @staticmethod
    def get_history(formal_contract_id: int) -> list[FormalContractDocument]:
        return FormalContractDocument.query.filter_by(
            formal_contract_id=formal_contract_id
        ).order_by(
            FormalContractDocument.generated_at.desc(),
            FormalContractDocument.id.desc(),
        ).all()
