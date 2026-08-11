"""Regression tests for the independent official-contract generator."""

from __future__ import annotations

import io
import json
from pathlib import Path

from docx import Document
from werkzeug.datastructures import FileStorage

from app import db
from app.models import (
    Company,
    Contract,
    FormalContract,
    FormalContractDocument,
    FormalContractParty,
    FormalContractSync,
    FormalContractTemplate,
    Product,
)
from app.services.official_contract_service import (
    OfficialContractService,
    OfficialContractValidationError,
)


def _template_file(tmp_path: Path) -> Path:
    """Build a small valid DOCX template with supported placeholders."""
    doc = Document()
    doc.add_paragraph('产品销售合同')
    doc.add_paragraph('合同编号：{{ contract.contract_no }}')
    doc.add_paragraph('甲方：{{ party_a.name }}')
    doc.add_paragraph('合计：{{ contract.total_amount_upper }}')
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = '品名'
    table.cell(0, 1).text = '数量'
    table.cell(0, 2).text = '金额'
    table.cell(1, 0).text = '{{ item.product_name }}'
    table.cell(1, 1).text = '{{ item.quantity }}'
    table.cell(1, 2).text = '{{ item.total_amount }}'
    path = tmp_path / 'formal_contract_template.docx'
    doc.save(path)
    return path


def _fixed_layout_template_file(tmp_path: Path) -> Path:
    """Build a minimal copy of the supplied fixed-layout contract template."""
    doc = Document()
    doc.add_paragraph('需方（甲方）：')
    doc.add_paragraph('供方（乙方）：江苏纯安科技有限公司')
    doc.add_paragraph('质量标准：')
    doc.add_paragraph('交货方式及费用承担：')
    doc.add_paragraph('交货时间及地点：')
    doc.add_paragraph('结算方式及期限：')
    doc.add_paragraph('违约责任：')
    doc.add_paragraph('解决合同纠纷的方式：')

    meta = doc.add_table(rows=3, cols=2)
    meta.cell(0, 0).text = '合同编号'
    meta.cell(0, 1).text = 'CA/'
    meta.cell(1, 0).text = '签订地点'
    meta.cell(1, 1).text = '南京'
    meta.cell(2, 0).text = '签订时间'
    meta.cell(2, 1).text = '年  月  日'

    items = doc.add_table(rows=3, cols=6)
    for index, header in enumerate(['序号', '品名', '规格型号', '数量', '单价(元)', '金额（元）']):
        items.cell(0, index).text = header
    items.cell(1, 0).text = '1'
    items.cell(2, 1).text = '合计： 元整'

    parties = doc.add_table(rows=1, cols=2)
    parties.cell(0, 0).text = '需      方 (甲方)'
    parties.cell(0, 1).text = '供     方（乙方）'

    path = tmp_path / 'fixed_layout_template.docx'
    doc.save(path)
    return path


def _seed_formal_contract(db_session, user_id: int):
    """Create a company, product, and formal contract through the service."""
    company = Company(name='模板测试甲方')
    product = Product(
        product_code='P-FORMAL-001',
        product_name='测试色谱柱',
        product_model='C18',
        product_type='自产',
        default_price=12.5,
    )
    db_session.add_all([company, product])
    db_session.commit()

    formal = OfficialContractService.save_formal_contract(
        {
            'party_a_name': company.name,
            'billing_address': '测试地址',
            'phone': '13800000000',
            'tax_no': 'TEST-TAX',
            'bank_name': '测试银行',
            'bank_account': '123456',
            'party_b_name': '测试乙方',
            'party_b_billing_address': '乙方地址',
            'party_b_phone': '13900000000',
            'party_b_tax_no': 'B-TAX',
            'party_b_bank_name': '乙方银行',
            'party_b_bank_account': '654321',
            'contract_no': 'CA/TEST-001',
            'sign_place': '南京',
            'sign_date': '2026-08-07',
            'quality_standard': '合格',
            'delivery_terms': '快递',
            'delivery_schedule': '七个工作日',
            'settlement_terms': '先款后货',
            'breach_terms': '依法处理',
            'dispute_terms': '协商解决',
        },
        [{
            'product_id': product.id,
            'quantity': '2',
            'unit_price': '12.5',
            'unit': '支',
        }],
        user_id,
    )
    return formal, company, product


def test_formal_contract_schema_and_party_defaults(app, db_session, base_data):
    """The new tables exist and selecting a party loads the latest contract number."""
    from sqlalchemy import inspect

    expected_tables = {
        'formal_contract_parties',
        'formal_contracts',
        'formal_contract_items',
        'formal_contract_templates',
        'formal_contract_documents',
        'formal_contract_syncs',
    }
    with app.app_context():
        assert expected_tables.issubset(set(inspect(db.engine).get_table_names()))
        formal, company, _ = _seed_formal_contract(
            db_session,
            base_data['superadmin_id'],
        )
        defaults = OfficialContractService.get_party_defaults(
            party_name=company.name,
        )
        assert defaults['contract_no'] == 'CA/TEST-001'
        assert defaults['billing_address'] == '测试地址'
        assert formal.total_amount == 25.0
        assert formal.party_a_billing_address == '测试地址'


def test_template_upload_activation_generation_and_snapshot(app, db_session, base_data, tmp_path):
    """DOCX generation uses the active template and stores an immutable snapshot."""
    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    template_path = _template_file(tmp_path)
    with app.app_context():
        upload = FileStorage(
            stream=io.BytesIO(template_path.read_bytes()),
            filename='formal_contract_template.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        template = OfficialContractService.create_template(
            upload,
            name='测试正式合同',
            version='v1.0',
            description='test',
            uploaded_by_id=base_data['superadmin_id'],
        )
        assert template.validation['unknown'] == []
        OfficialContractService.activate_template(template.id)
        document = OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        assert Path(document.docx_path).is_file()
        assert document.template_version == 'v1.0'
        snapshot = json.loads(document.snapshot_json)
        assert snapshot['contract']['contract_no'] == 'CA/TEST-001'
        assert snapshot['contract']['sign_date_display'] == '2026年08月07日'
        assert snapshot['items'][0]['product_name'] == '测试色谱柱'

        generated = Document(document.docx_path)
        text = '\n'.join(p.text for p in generated.paragraphs)
        assert 'CA/TEST-001' in text
        assert '模板测试甲方' in text


def test_fixed_layout_template_is_filled(app, db_session, base_data, tmp_path):
    """The supplied static template is filled without requiring Jinja tags."""
    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    template_path = _fixed_layout_template_file(tmp_path)
    with app.app_context():
        upload = FileStorage(
            stream=io.BytesIO(template_path.read_bytes()),
            filename='fixed_layout_template.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        template = OfficialContractService.create_template(
            upload,
            name='固定版式模板',
            version='v1',
            description=None,
            uploaded_by_id=base_data['superadmin_id'],
        )
        assert template.validation['mode'] == 'fixed_layout'
        OfficialContractService.activate_template(template.id)
        document = OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )

        generated = Document(document.docx_path)
        all_text = '\n'.join(
            [paragraph.text for paragraph in generated.paragraphs]
            + [
                cell.text
                for table in generated.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        assert 'CA/TEST-001' in all_text
        assert '模板测试甲方' in all_text
        assert '测试乙方' in all_text
        assert '测试色谱柱' in all_text
        assert '￥25.00' in all_text


def test_formal_contract_party_fields_are_snapshotted(app, db_session, base_data):
    """Later party-profile edits do not rewrite an older formal contract."""
    first, company, product = _seed_formal_contract(
        db_session,
        base_data['superadmin_id'],
    )
    with app.app_context():
        second = OfficialContractService.save_formal_contract(
            {
                'party_a_name': company.name,
                'billing_address': '第二次地址',
                'phone': '13999999999',
                'tax_no': 'SECOND-TAX',
                'bank_name': '第二次银行',
                'bank_account': 'SECOND-ACCOUNT',
                'party_b_name': '测试乙方',
                'contract_no': 'CA/TEST-002',
                'sign_place': '南京',
                'sign_date': '2026-08-07',
            },
            [{
                'product_id': product.id,
                'quantity': '1',
                'unit_price': '12.5',
                'unit': '支',
            }],
            base_data['superadmin_id'],
        )
        first_context = OfficialContractService.build_context(
            FormalContract.query.get(first.id),
        )
        second_context = OfficialContractService.build_context(second)
        assert first_context['party_a']['billing_address'] == '测试地址'
        assert second_context['party_a']['billing_address'] == '第二次地址'


def test_form_route_creates_new_party_and_reloads_defaults(
    app,
    client,
    login,
    db_session,
    base_data,
):
    """The browser form closes the new-party and recent-defaults loop."""
    product = Product(
        product_code='P-ROUTE-001',
        product_name='路由测试产品',
        product_model='M-ROUTE',
        product_type='自产',
        default_price=8.8,
    )
    db_session.add(product)
    db_session.commit()
    login(base_data['superadmin_id'])

    response = client.post(
        '/official-contract/new',
        data={
            'party_a_name': '路由新甲方',
            'billing_address': '路由地址',
            'phone': '13200000000',
            'tax_no': 'ROUTE-TAX',
            'bank_name': '路由银行',
            'bank_account': 'ROUTE-ACCOUNT',
            'party_b_name': '路由乙方',
            'party_b_billing_address': '乙方地址',
            'party_b_phone': '13300000000',
            'party_b_tax_no': 'ROUTE-B-TAX',
            'party_b_bank_name': '乙方银行',
            'party_b_bank_account': 'ROUTE-B-ACCOUNT',
            'contract_no': 'ROUTE-FORMAL-001',
            'sign_place': '南京',
            'sign_date': '2026-08-07',
            'item_count': '1',
            'item_0_product_id': str(product.id),
            'item_0_product_code': product.product_code,
            'item_0_product_name': product.product_name,
            'item_0_product_model': product.product_model,
            'item_0_quantity': '3',
            'item_0_unit': '支',
            'item_0_unit_price': '8.8',
        },
        follow_redirects=False,
    )
    assert response.status_code == 302
    assert '/official-contract/' in response.headers['Location']

    with app.app_context():
        party = FormalContractParty.query.filter_by(
            party_a_name='路由新甲方',
        ).first()
        assert party is not None
        formal = FormalContract.query.filter_by(
            party_id=party.id,
            contract_no='ROUTE-FORMAL-001',
        ).first()
        assert formal is not None
        assert formal.total_amount == 26.4
        party_id = party.id
        formal_id = formal.id

    defaults_response = client.get(
        f'/official-contract/api/parties/defaults?party_id={party_id}',
    )
    assert defaults_response.status_code == 200
    defaults = defaults_response.get_json()
    assert defaults['contract_no'] == 'ROUTE-FORMAL-001'
    assert defaults['billing_address'] == '路由地址'

    edit_response = client.get(f'/official-contract/{formal_id}/edit')
    assert edit_response.status_code == 200
    assert 'party_b_billing_address' in edit_response.get_data(as_text=True)


def test_formal_contract_print_route_and_download(app, client, login, db_session, base_data, tmp_path):
    """The browser print page follows the existing window.print flow."""
    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    template_path = _template_file(tmp_path)
    with app.app_context():
        upload = FileStorage(
            stream=io.BytesIO(template_path.read_bytes()),
            filename='print_template.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        template = OfficialContractService.create_template(
            upload,
            name='打印模板',
            version='v1',
            description=None,
            uploaded_by_id=base_data['superadmin_id'],
        )
        OfficialContractService.activate_template(template.id)

    login(base_data['superadmin_id'])
    print_response = client.get(
        f'/official-contract/{formal.id}/print?autoprint=0',
        follow_redirects=False,
    )
    assert print_response.status_code == 200
    html = print_response.get_data(as_text=True)
    assert 'window.print()' in html
    assert 'CA/TEST-001' in html

    with app.app_context():
        document = FormalContractDocument.query.filter_by(
            formal_contract_id=formal.id,
        ).first()
        download_response = client.get(
            f'/official-contract/{formal.id}/documents/{document.id}/download',
        )
        assert download_response.status_code == 200
        assert download_response.mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def test_print_route_uses_generated_document_snapshot(
    app,
    client,
    login,
    db_session,
    base_data,
    tmp_path,
):
    """Printing an old file keeps the exact data saved when it was generated."""
    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    template_path = _template_file(tmp_path)
    with app.app_context():
        upload = FileStorage(
            stream=io.BytesIO(template_path.read_bytes()),
            filename='snapshot_print_template.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        template = OfficialContractService.create_template(
            upload,
            name='snapshot-print-template',
            version='v1',
            description=None,
            uploaded_by_id=base_data['superadmin_id'],
        )
        OfficialContractService.activate_template(template.id)
        document = OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        formal_id = formal.id
        document_id = document.id
        formal.contract_no = 'CA/CHANGED-AFTER-GENERATION'
        formal.party_b_name = 'Changed Party B'
        db_session.commit()

    login(base_data['superadmin_id'])
    response = client.get(
        f'/official-contract/{formal_id}/print?document_id={document_id}&autoprint=0',
    )
    assert response.status_code == 200
    html = response.get_data(as_text=True)
    assert 'CA/TEST-001' in html
    assert '测试乙方' in html
    assert 'CA/CHANGED-AFTER-GENERATION' not in html
    assert 'Changed Party B' not in html
    assert '2026年08月07日' in html
    assert 'margin: 21.98mm 13.16mm 11.01mm 24.01mm' in html
    assert 'width: 65.88mm' in html
    assert 'width: 167.83mm' in html
    assert 'body > main.container' in html
    assert 'body > footer' in html
    assert 'min-height: 0;' in html


def test_sync_to_transaction_contract_is_idempotent(app, db_session, base_data, tmp_path):
    """Sync creates one ERP contract and repeated calls return the same link."""
    formal, company, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    template_path = _template_file(tmp_path)
    with app.app_context():
        upload = FileStorage(
            stream=io.BytesIO(template_path.read_bytes()),
            filename='sync_template.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        template = OfficialContractService.create_template(
            upload,
            name='同步模板',
            version='v1',
            description=None,
            uploaded_by_id=base_data['superadmin_id'],
        )
        OfficialContractService.activate_template(template.id)
        OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )

        first = OfficialContractService.sync_to_transaction_contract(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        second = OfficialContractService.sync_to_transaction_contract(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        assert first.id == second.id
        assert Contract.query.filter_by(company_name=company.name).count() == 1
        assert FormalContractSync.query.filter_by(formal_contract_id=formal.id).count() == 1
        fresh_formal = FormalContract.query.get(formal.id)
        assert fresh_formal.status == 'synced'


def test_template_manager_and_limited_user_permissions(client, login, base_data):
    """A user without formal-contract permission cannot open the module."""
    login(base_data['limited_user_id'])
    response = client.get('/official-contract/', follow_redirects=False)
    assert response.status_code == 302
    assert '/erp/' in response.headers['Location']


def test_builtin_template_is_available_without_manual_upload(
    app,
    db_session,
    base_data,
):
    """The bundled DOCX is registered as the global fallback automatically."""
    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    with app.app_context():
        template = OfficialContractService.get_active_template()
        assert template is not None
        assert template.department_id is None
        assert template.version == 'builtin-v1'
        assert template.is_active
        assert Path(template.stored_path).is_file()

        document = OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        assert Path(document.docx_path).is_file()
        assert document.template_id == template.id


def test_department_templates_are_isolated_and_override_global_fallback(
    app,
    db_session,
    base_data,
    tmp_path,
):
    """Activating one department template must not deactivate another department."""
    from app.models import Department

    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    department_a = Department(name='Assembly')
    department_b = Department(name='Research')
    db_session.add_all([department_a, department_b])
    db_session.commit()
    formal.department_id = department_a.id
    db_session.commit()

    template_path = _template_file(tmp_path)
    with app.app_context():
        def upload(name: str, department_id: int):
            return OfficialContractService.create_template(
                FileStorage(
                    stream=io.BytesIO(template_path.read_bytes()),
                    filename=f'{name}.docx',
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                ),
                name=name,
                version='v1',
                description=None,
                uploaded_by_id=base_data['superadmin_id'],
                department_id=department_id,
            )

        assembly_template = upload('assembly-template', department_a.id)
        research_template = upload('research-template', department_b.id)
        OfficialContractService.activate_template(assembly_template.id)
        OfficialContractService.activate_template(research_template.id)

        assert OfficialContractService.get_active_template(department_a.id).id == assembly_template.id
        assert OfficialContractService.get_active_template(department_b.id).id == research_template.id
        assert FormalContractTemplate.query.get(assembly_template.id).is_active
        document = OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        assert document.template_id == assembly_template.id


def test_form_contains_department_and_product_library_dropdown(
    app,
    client,
    login,
    db_session,
    base_data,
):
    """The browser form exposes both department and product-library selectors."""
    product = Product(
        product_code='P-DROPDOWN-001',
        product_name='Dropdown Test Product',
        product_model='M-DROPDOWN',
        default_price=10,
    )
    db_session.add(product)
    db_session.commit()
    login(base_data['superadmin_id'])

    response = client.get('/official-contract/new')
    assert response.status_code == 200
    html = response.get_data(as_text=True)
    assert 'name="department_id"' in html
    assert 'class="form-select product-select' in html
    assert 'P-DROPDOWN-001' in html


def test_template_upload_route_persists_department_scope(
    app,
    client,
    login,
    base_data,
    tmp_path,
):
    """The admin upload route passes the selected department to persistence."""
    template_path = _template_file(tmp_path)
    login(base_data['superadmin_id'])
    response = client.post(
        '/official-contract/templates/upload',
        data={
            'name': 'department-route-template',
            'version': 'v1',
            'department_id': str(base_data['department_id']),
            'description': 'route test',
            'template_file': (
                io.BytesIO(template_path.read_bytes()),
                'department-route-template.docx',
            ),
        },
        content_type='multipart/form-data',
        follow_redirects=False,
    )
    assert response.status_code == 302
    with app.app_context():
        template = FormalContractTemplate.query.filter_by(
            name='department-route-template',
        ).one()
        assert template.department_id == base_data['department_id']


def test_builtin_fixed_layout_preserves_template_geometry(
    app,
    db_session,
    base_data,
):
    """Filling the supplied DOCX must not rebuild its page or table geometry."""
    formal, _, _ = _seed_formal_contract(db_session, base_data['superadmin_id'])
    source = Path(__file__).resolve().parents[1] / (
        'templates/official_contract/builtin/配件合同模版.docx'
    )
    with app.app_context():
        template = OfficialContractService.create_template(
            FileStorage(
                stream=io.BytesIO(source.read_bytes()),
                filename='supplied-layout.docx',
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            ),
            name='supplied-layout',
            version='v1',
            description=None,
            uploaded_by_id=base_data['superadmin_id'],
        )
        OfficialContractService.activate_template(template.id)
        document = OfficialContractService.generate_document(
            formal.id,
            user_id=base_data['superadmin_id'],
        )
        from docx import Document

        original = Document(source)
        generated = Document(document.docx_path)
        assert len(original.sections) == len(generated.sections)
        for original_section, generated_section in zip(
            original.sections,
            generated.sections,
        ):
            assert (
                original_section.page_width,
                original_section.page_height,
                original_section.top_margin,
                original_section.right_margin,
                original_section.bottom_margin,
                original_section.left_margin,
            ) == (
                generated_section.page_width,
                generated_section.page_height,
                generated_section.top_margin,
                generated_section.right_margin,
                generated_section.bottom_margin,
                generated_section.left_margin,
            )

        assert len(original.tables) == len(generated.tables)
        for original_table, generated_table in zip(
            original.tables,
            generated.tables,
        ):
            assert original_table._tbl.tblPr.xml == generated_table._tbl.tblPr.xml
            assert original_table._tbl.tblGrid.xml == generated_table._tbl.tblGrid.xml
            assert len(original_table.rows) == len(generated_table.rows)
            for original_row, generated_row in zip(
                original_table.rows,
                generated_table.rows,
            ):
                for original_cell, generated_cell in zip(
                    original_row.cells,
                    generated_row.cells,
                ):
                    assert original_cell._tc.tcPr.xml == generated_cell._tc.tcPr.xml
